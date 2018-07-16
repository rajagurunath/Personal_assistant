# -*- coding: utf-8 -*-
"""
Created on Sun Jan 21 10:54:35 2018

@author: welcome
"""
from __future__ import absolute_import
from __future__ import division, print_function, unicode_literals

import requests
from lxml import html

import easygui
import win32com.client as wincl
speak = wincl.Dispatch("SAPI.SpVoice")
import dateutil.parser                                                      
import os
import pandas as pd
import seaborn as sns
import matplotlib
#matplotlib.use('Agg')
import matplotlib.pyplot as plt
import glob
from win32com.client import Dispatch
import datetime
from bs4 import BeautifulSoup
from sumy.parsers.html import HtmlParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer as Summarizer
from sumy.nlp.stemmers import Stemmer
from sumy.utils import get_stop_words
from dateutil.parser import parse
plt.style.use('fivethirtyeight')
from docx import Document

import win32com.client as wincl
speak = wincl.Dispatch("SAPI.SpVoice")
from pytrends.request import TrendReq
from arivx import query
try:
	from urllib.parse import quote
except ImportError:
	from urllib import quote
from functools import reduce

cdir=os.getcwd()
mail_dir=cdir+'\Mail_documents'

amazondir=cdir+'\Amazon_purchase'

sumydir=cdir+'\summarization'
popdir=cdir+'\Popularity_analysis'
jourdir=cdir+'\Journals_abstract'

if not os.path.exists(mail_dir):
    os.mkdir(mail_dir)
if not os.path.exists(amazondir):
    os.mkdir(amazondir)
if not os.path.exists(sumydir):
    os.mkdir(sumydir)
if not os.path.exists(popdir):
    os.mkdir(popdir)
if not os.path.exists(jourdir):
    os.mkdir(jourdir)

os.chmod(mail_dir,777)
        
os.chmod(mail_dir,777)
os.chmod(amazondir,777)
os.chmod(sumydir,777)
os.chmod(popdir,777)
os.chmod(jourdir,777)

SITE_URL = 'https://www.amazon.com'
BASE_URL = '%s/s/ref=sr_qz_back?sf=qz&keywords=%s&ie=UTF8&tag=alhsabc-20&unfiltered=1&page=%s' % (SITE_URL,'%s','%s')
OUT_URL_TAG = '&tag=alhsabc-20'
OUT_URL_REF = '/ref=as_li_tl?tag=alhsabc-20'
URL_HEADERS = {'User-Agent' : 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_11_3) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/60.0.3112.113 Safari/537.36'}
DEFAULT_PRICE_TEXT = "Base price"
MAX_COL_SIZE = 80

categories = ['cs.', 'stat.', 'q-bio.', 'nlin.', 'math.',
              'astro-ph', 'cond-mat.', 'gr-qc', 'hep-ex',
              'hep-lat', 'hep-ph', 'hep-th', 'math-ph', 'nucl-ex',
              'nucl-th', 'physics.', 'quant-ph']





def timesheet_with_graph():
    try:
        filepath=timesheet()
  
        df=pd.read_csv(filepath,encoding = "ISO-8859-1")

        speak.Speak("Please enter filename to save:")
        file_name=easygui.textbox(msg='Please enter filename to save:',title='Time_sheet_with_graph').split()[0]
        
        
        if not os.path.exists(mail_dir):
            os.mkdir(mail_dir)
        os.chmod(mail_dir,777)
            
        df.set_index(df['Title'],inplace=True)
    #    os.chmod(current_dir+'\\'+str(file_name)+'\\',777)
        fig, ax = plt.subplots()
    # the size of A4 paper
        fig.set_size_inches(10, 20)
        sns.barplot(df['Duration(Minutes)'],df.index,data=df)
        plt.savefig(mail_dir+'\\'+str(file_name)+'.png',format='png')
#        print('file saved to '+mail_dir+'\\'+str(file_name)+'.png'))
        easygui.msgbox(msg='file saved to '+str(mail_dir)+'\\'+str(file_name)+'.png',title='Timesheet with graph prepared')
        speak.Speak('file saved to directory' +str(file_name)+'with the name time spending in meeting')
    except Exception as e:
        speak.Speak('Sorry My mistake please provide your feedback regarding this error')
        easygui.exceptionbox(str(e))
        
        
def timesheet_with_graph_for_specificDay():
    try:
        filepath=timesheet()
        speak.Speak("Please Enter the date for which you want your analysis")
    
        date=easygui.textbox(msg='Enter the date (dd-mm-yyyy):',title='Time sheet with graph')
    #    filepath='D:\karparthy\yy.csv'
    
    #    sns.set_style('whitegrid')
        current_dir=os.getcwd()
        filepath=r''+filepath
        print(filepath)
        date=parse(date).date()
        df=pd.read_csv(filepath,encoding = "ISO-8859-1")
    #    df.columns
        speak.Speak("Please Enter the filename to save")
    
        file_name=easygui.textbox(msg='Please enter filename to save:',title='Time_sheet_with_graph').split()[0]
        
    #    file_name=input('Enter name of the directory to save')
        
        if not os.path.exists(current_dir+'\\'+str(file_name)):
            os.mkdir(current_dir+'\\'+str(file_name))
            
    
        df['Start_date']=pd.DatetimeIndex(df['Start']).date
        df.set_index(df['Title'],inplace=True)
        df=df[df['Start_date']==date]
    #    os.chmod(current_dir+'\\'+str(file_name)+'\\',777)
        
    #    df['Start']=
        
    #    df=df[df['Start']==]
        fig, ax = plt.subplots()
    # the size of A4 paper
        fig.set_size_inches(10, 5)
        sns.barplot(df['Duration(Minutes)'],df.index,data=df)
        plt.savefig(current_dir+'\\'+str(file_name)+'\\'+str('Time_spent_meeting')+'.png',format='png')
        print('file saved to '+str(current_dir+'\\'+str(file_name)+'\\'+str('Time_spent_meeting')+'.png'))
        speak.Speak("File saved master")
    
        easygui.msgbox(msg='file saved to '+str(current_dir+'\\'+str(file_name)+'\\'+str('Time_spent_meeting')+'.png'),title='Timesheet with graph prepared')
    except Exception as e:
        speak.Speak('Sorry My mistake please provide your feedback regarding this error')
        easygui.exceptionbox(str(e))
        
    



def timesheet():
    try:
        OUTLOOK_FORMAT = '%m/%d/%Y %H:%M'
        outlook = Dispatch("Outlook.Application")
        ns = outlook.GetNamespace("MAPI")
    
        appointments = ns.GetDefaultFolder(9).Items 
    
    # Restrict to items in the next 30 days (using Python 3.3 - might be slightly different for 2.7)
        begin = datetime.date.today()
        end = begin + datetime.timedelta(days = 30);
        restriction = "[Start] > '" + begin.strftime("%m/%d/%Y") + "' AND [End] < '" +end.strftime("%m/%d/%Y") + "'"
    #    restrictedItems = appointments.Restrict(restriction)
    
        #appointments.Sort("[Duration]")
        appointments.IncludeRecurrences = "True"
    
        # Iterate through restricted AppointmentItems and print them
        calcTableHeader = ['Title', 'Organizer', 'Start', 'Duration(Minutes)'];
        calcTableBody = [];
    
        #pdb.set_trace()
        for appointmentItem in appointments:
            row = []
            row.append(appointmentItem.Subject)
            row.append(appointmentItem.Organizer)
            row.append(appointmentItem.Start.Format(OUTLOOK_FORMAT))
            row.append(appointmentItem.Duration)
            calcTableBody.append(row)
    
        df=pd.DataFrame(calcTableBody, columns=calcTableHeader)
        dir_=os.getcwd()
        speak.Speak("Please Enter the filename to save")
    
        file_name=easygui.textbox(msg='Please enter filename to save:',title='Time_sheet').split()[0]
    #    file_name = input("Please enter filename to save: ")
    #    resultsfile = open(dir_+'\\'+str(file_name)+'.csv','w',encoding="utf-8")
#        if not os.path.exists(mail_dir+'\\'+str(file_name)):
#            os.mkdir(mail_dir+'\\'+str(file_name))
        
        filepath=mail_dir+'\\'+str(file_name)+'.csv'
        df.to_csv(filepath)
    #    print('your Timesheet saved '+str(filepath))
        speak.Speak("your timesheet saved in the following path")
    
        easygui.msgbox(msg='your Timesheet saved '+str(filepath),title='Timesheet prepared')
        return filepath
    except Exception as e:
        speak.Speak('Sorry My mistake please provide your feedback regarding this error')
        easygui.exceptionbox(str(e))
        


def getTextFromURL(url):
    r = requests.get(url)
    soup = BeautifulSoup(r.text, "html.parser")
    title = ' '.join(map(lambda p: p.text, soup.find_all('h1')))
    return title



def summarize(SENTENCES_COUNT):
    try:
        #    url = "https://machinelearningmastery.com/prepare-text-data-machine-learning-scikit-learn/"
        speak.Speak("Please Enter the U r l for summarization")
    
        url=easygui.textbox(msg='Enter url for which you want summarization:',title='Summarization').split()[0]
        title=getTextFromURL(url)
    #    url="https://medium.com/@gurunathrajagopal/what-happens-when-machines-are-curious-to-learn-9aed6805bf36"
        parser = HtmlParser.from_url(url, Tokenizer(LANGUAGE))
        # or for plain text files
        # parser = PlaintextParser.from_file("document.txt", Tokenizer(LANGUAGE))
        stemmer = Stemmer(LANGUAGE)
    
        summarizer = Summarizer(stemmer)
        summarizer.stop_words = get_stop_words(LANGUAGE)
        string_dict={}
        for idx,sentence in enumerate(summarizer(parser.document, SENTENCES_COUNT)):
    #        f.write(str(sentence))
            string_dict[idx]=str(sentence)
    #        print(type(sentence))
    #    print(string_dict)
#        speak.Speak("Please Enter the filename to save to summarization")
    
#        file_name=easygui.textbox(msg='Enter filename to save the summarization:',title='Summarization').split()[0]
#        current_dir=os.getcwd()
    #    f=open(current_dir+'\\'+str(file_name)+'.txt','w') 
    #    f.write('Summarization')
        document = Document()
    
        document.add_heading('Summarization of '+str(title), 0)
        p=document.add_paragraph('Summarizing your article in crisp {} points'.format(SENTENCES_COUNT))
    
        for idx,sent in zip(string_dict.keys(),string_dict.values()):
            adding_break=p.add_run()
            adding_break.add_break()
            p=document.add_paragraph(sent)
        adding_break=p.add_run()
        adding_break.add_break()    
        document.save(sumydir+'\\'+'summarization.docx')    
        speak.Speak("Summarization was saved to the following path")
    
       
    #        f.write('\n')
    #        f.write(str(idx))
    #        f.write('.  ')
    #        f.write(sent)
    #    f.close()
        easygui.msgbox(msg='Sumarized file saved in this file '+sumydir+'\\'+'summarization.docx',title='Summarization')
    except Exception as e:
        speak.Speak('Sorry My mistake please provide your feedback regarding this error')
        easygui.exceptionbox(str(e))
def google_popularity():
    try:
        pytrend = TrendReq()
        
        # Create payload and capture API tokens. Only needed for interest_over_time(), interest_by_region() & related_queries()
        cap_list=['Compare search metrics','Trending searches']
        speak.Speak("Welcome to the Google search capabilities")
        
        google_response_=easygui.choicebox(msg='Google search capabilities',title='Google Trends',choices=cap_list)
        print(google_response_)
        if google_response_=='Compare search metrics':
            
            Compare_two_keywords(pytrend)
        elif google_response_=='Trending searches':
            trending_searches_df = pytrend.trending_searches()
            trending_searches_df.to_csv(popdir+'\googleTrendingSearch'+str(datetime.datetime.now().date())+'.csv')
            easygui.msgbox(msg='Saving latest google trending searches in this csv'+os.getcwd()+os.getcwd()+'\googleTrendingSearch'+str(datetime.datetime.now().date())+'.csv')
#            print(trending_searches_df.head())

    except Exception as e:
        speak.Speak('Sorry My mistake please provide your feedback regarding this error')
        easygui.exceptionbox(str(e))
        
        
#print(type(google_response_))
#easygui.msgbox(msg='Google search capabilities',title='Google Trends')
#easygui.textbox(msg='Enter url for which you want summarization:',title='Summarization')

#pytrend = TrendReq()

def Compare_two_keywords(pytrend):
    directory=popdir
    directory=directory+'\\plots'

    number =int(easygui.textbox(msg='How many keywords you want to compare ?'))
    kw_list=[]
    for i in range(1,number+1):
        search_term=easygui.textbox(msg='Enter the Keyword '+str(i),title='Comparing  search terms')
        kw_list.append(str(search_term))
#        search_term_2=easygui.textbox(msg='Enter the Keyword 2',title='Comparing two search terms')
    
    pytrend.build_payload(kw_list=kw_list)
    document = Document()

    document.add_heading('Visualizing the popularity of the keywords by using no of google search per keyword', 0)
    p=document.add_paragraph('No of Keywords to be compared '+str(number))

    
# Interest Over Time
    interest_over_time_df = pytrend.interest_over_time()
    for i in range(0,number):
        plt.plot(interest_over_time_df.index,interest_over_time_df[kw_list[i]],label=str(kw_list[i]))
#        sns.tsplot(data=interest_over_time_df[str(kw_list[i])],legend=True)
        plt.xlabel('Time')
        plt.title('comparison of search terms')
        plt.legend()
#        plt.
#        sns.tsplot(data=interest_over_time_df[str(kw_list[]])
#    plt.save
    if not os.path.exists(directory):
        os.mkdir(directory)
    plt.tight_layout()
    plt.savefig(directory+'\interest_over_time.png')
    plt.close()
    adding_break=p.add_run()
    adding_break.add_break()    
    document.add_paragraph('visualizing the popularity')    
    adding_break=p.add_run()
    adding_break.add_break()    
    document.add_picture(directory+'\interest_over_time.png')

    
    
    interest_by_region_df = pytrend.interest_by_region()
#    print(interest_by_region_df)
    for i in range(0,number):
        print(1)
        Top_10_search_term=interest_by_region_df.sort_values(by=str(kw_list[i]))[:10]
        print(Top_10_search_term)
        sns.barplot(Top_10_search_term[str(kw_list[i])],Top_10_search_term.index)
#        plt.show()
#        print(1)
        plt.tight_layout()
#        plt.show()
        print(directory+'\\'+'interest_by_region_'+str(kw_list[i])+'.png')
        plt.savefig(directory+'\\'+'interest_by_region_'+str(kw_list[i])+'.png')
        plt.close()
        document.add_paragraph('visualizing the popularity by region')    
        document.add_picture(directory+'\\'+'interest_by_region_'+str(kw_list[i])+'.png')
#    p.add_paragraph('visualizing the popularity')    
    adding_break=p.add_run()
    adding_break.add_break()    
    document.save(popdir+'\popularity_visualize.docx')
    



def getItem(query_string,page_num,item_num):

	(products,url) = getSearchPage(query_string,page_num=page_num)
	try:
		item = products[item_num]
	except KeyError:
		raise ValueError('The item number %s could not be found on page %s' % (item_num,page_num))

	url = item['url']
	if 'ref=' in url:
		url += OUT_URL_TAG
	else:
		url += OUT_URL_REF

	return ({item_num:item},url)


def getSearchPage(query_string,page_num=1):
    contentlist=[],[],[]
    for i in range(1,6):
         
        (cont,url) = getHtmlUrl(query_string,page_num=i)
        contentlist.append(cont)
    products = getProducts(contentlist)
    df=pd.DataFrame(products).T
    directory=amazondir
    df.to_csv(directory+'\\'+str(query_string)+'_amazon.csv')
    speak.Speak('your product details are prepared, Happy purchasing')
    easygui.msgbox(msg='directory'+'\\'+str(query_string)+'_amazon.csv',title='Amazon-purchase')   
    return directory+'\\'+str(query_string)+'_amazon.csv'
#    urllist.append(url)
#    productlist.append(products)
#    return (productlist,urllist)

#def frame(tuple_)->list:
#    productlist=tuple_[0]
#    urllist=tuple_[0]
#    tmplist=[]
#    for pdict in productlist:
#        tmpdf=pd.DataFrame(pdict)
#        tmplist.append(tmpdf)
#    dataframe=pd.concat(tmplist)
#    return dataframe

def getProducts(contentlist):
    products = {}
    for content in contentlist:
    	tree = html.fromstring(content)

    
    	results = tree.xpath('//li[contains(@id,"result_")]')
    
    	for res in results:
    		num = res.xpath('./@id')[0]
    		num = str(num)
    		num = num.rsplit('_',1)[-1]
    		num = int(num)
    
    		name = res.xpath('.//*[self::h2 or self::span]/@data-attribute')
    		url = res.xpath('.//a[h2]/@href')
    
    		if len(name) > 0 and len(url) > 0:
    			name = str(name[0])	
    			url = str(url[0])
    
    			if url.startswith('/'):
    				url = SITE_URL + url
    
    			rating = [x.text for x in res.xpath('.//*[contains(@class,"a-icon-star")]/span')]
    
    			rows = res.xpath('.//div[@class="a-column a-span7" or @class="s-item-container"]/div[contains(@class,"a-row")]')
    			temp_title = DEFAULT_PRICE_TEXT
    			prices = {}
    			for row in rows:
    				get_title = [x.text for x in row.xpath('.//h3')]
    				get_prices = row.xpath('.//span/@aria-label')
    	
    				if len(get_title) > 0:
    					temp_title = get_title[0]
    				elif len(get_prices) > 0:
    					prices[temp_title] = get_prices
    					temp_title = DEFAULT_PRICE_TEXT
    	
    			products[num] = {
    				'name' : name,
    				'url' : url,
    				'prices' : prices,
    				'rating' : rating[0] if len(rating) > 0 else ''
    			}
    	
    return products

def getHtmlUrl(query_string,page_num=1):
	url = BASE_URL % (quote(query_string),page_num)
	req = requests.get(url,headers=URL_HEADERS)

	if not req.ok:
		raise ValueError('The requested page could not be found')

	return (req.content.decode('utf8', errors='ignore'),url)

def getCleanPrices(price_dict):

	price_texts = reduce(lambda a, b: a + b[1],price_dict.items(),[])
	price_texts = reduce(lambda x, y: x + y.split('-'),price_texts,[])

	out_prices = []
	for text in price_texts:
		val = float(re.sub('[^\d.]','',text))
		out_prices.append(val)



	return sorted(out_prices)

def getRatingValue(text):
	val = text.split()

	try:
		int_val = float(val[0])
		return int_val
	except:
		return 0
def amazon_products():
    try:
        query_string=str(easygui.textbox(msg='Enter the product name ',title='Getting Amazon products '))
        getSearchPage(query_string)
    except Exception as e:
        speak.Speak('Sorry My mistake please provide your feedback regarding this error')
        easygui.exceptionbox(msg=str(e),title='sorry for the inconvenince')

        

def journal_downloading():
    try:
#        import arxivpy
        #import easygui
        #from docx import Document
        #from docx.shared import Inches
        #import os
        
        topic =str(easygui.textbox(msg='Enter the topic you want the paper ?'))
        print(topic)
        max_ind=10
        articles=query(search_query=topic,
                                 start_index=0, max_index=max_ind, results_per_iteration=100,
                                 wait_time=5.0, sort_by='lastUpdatedDate') # grab 200 articles
        
        print(articles)
        
        document = Document()
#        directory=os.getcwd()
#        directory=directory+'\Journals_abstract'
        if not os.path.exists(jourdir):
                os.mkdir(jourdir)
            
        
        document.add_heading(topic, 0)
        for dictionary in articles:
            document.add_heading(dictionary.get('title'),0)
            p=document.add_paragraph(dictionary.get('abstract'))
            adding_break=p.add_run()
            adding_break.add_break()
            p.add_run(' main_author: '+str(dictionary.get('main_author'))).bold=True
            adding_break=p.add_run()
            adding_break.add_break()
            p.add_run(' published_date: '+str(dictionary.get('publish_date'))).bold=True
            adding_break=p.add_run()
        
            adding_break.add_break()
            p.add_run(' pdf_url: '+str(dictionary.get('pdf_url'))).italic=True
            adding_break=p.add_run()
        
            adding_break.add_break()
        document.save(jourdir+'\\'+str(topic)+'_abstract.docx')    
    except Exception as e:
        speak.Speak('Sorry My mistake please provide your feedback regarding this error')
        easygui.exceptionbox(msg=str(e),title='sorry for the inconvenince')


def feedback_mail():
    
    easygui.msgbox('your message will be send to XXXXX,XXXXX@company_name.com')
    try:
        import smtplib
    
    #from string import Template
    
        from email.mime.multipart import MIMEMultipart
        from email.mime.text import MIMEText
#        import easygui
        crendlist=easygui.multpasswordbox(msg='outlook crendentials',fields=['Username','Password'])
        MY_ADDRESS =crendlist[0]
        PASSWORD = crendlist[1]
        easygui.msgbox('your message will be send to XXXXX,XXXXX@company_name.com')

        name, email = 'XXXXX','XXXXX@outlook.com'
    #    message_template = read_template('message.txt')
    
        # set up the SMTP server
        s = smtplib.SMTP(host='outlook.company_name.com', port=587)
    
    #    s = smtplib.SMTP(host='your_host_address_here', port=your_port_here)
        s.starttls()
        s.login(MY_ADDRESS, PASSWORD)
    
        # For each contact, send the email:
    #    for name, email in zip(names, emails):
        msg = MIMEMultipart()       # create a message
    #
    #        # add in the actual person name to the message template
    #        message = message_template.substitute(PERSON_NAME=name.title())
    #
    #       # Prints out the message body for our sake
        speak.Speak('Enter your feedback message')
        
        message=str(easygui.textbox(msg='Enter your feedback',title='feedback'))
        print(message)
        
            # setup the parameters of the message
        msg['From']=MY_ADDRESS
        msg['To']=email
        msg['Subject']="Feedback for personal assistant"
            
            # add in the message body
        msg.attach(MIMEText(message, 'plain'))
            
            # send the message via the server set up earlier.
        s.send_message(msg)
        del msg
            
        # Terminate the SMTP session and close the connection
        s.quit()
    except Exception as e:
        speak.Speak('Sorry My mistake please provide your feedback regarding this error')
        easygui.exceptionbox(str(e))



def get_mail(type_):
    if type_=='sentbox':
        folder_id=5
        mail_df=pd.DataFrame(columns=['SentTo','subject','senttime','body','date'])
        

    elif type_=='inbox':
        folder_id=6
        mail_df=pd.DataFrame(columns=['sender_name','subject','receivedtime','body','date'])

#    sent_item=5
#    inbox=6
    import win32com.client 
    outlook = win32com.client.Dispatch("Outlook.Application").GetNamespace("MAPI")
    inbox = outlook.GetDefaultFolder(folder_id) # "6" refers to the index of a folder - in this case,
                                            # the inbox. You can change that number to reference
                                            # any other folder
    messages = inbox.Items
    message = messages.GetFirst()
#    rec_time = message.CreationTime
#    body_content = message.body
#    subj_line = message.subject
    msglst=[]
    msgbdy=[]
    msgtme=[]
    sender_list=[]
#    body_list=[]
    while message:
        try:
    #        print(message.subject, message.CreationTime)
            
            msglst.append(message.subject)
            msgtme.append(message.CreationTime)
            msgbdy.append(message.body)
            sender =  message.SenderName
            sender_list.append(sender)
            message = messages.GetNext()
        except :
            message = messages.GetNext()
#    import pandas as pd
    
    #now_datetime = datetime.datetime.fromtimestamp (int (msgtme[0]))
    
    
    
    datelst=[dateutil.parser.parse(str(msgtme[i])) for i in range(0,len(msgtme))]
    #ss=[datetime.datetime.strptime(str(datelst[i]).rstrip("+00:00"), '%Y-%m-%d %H:%M:%S') for i in range(len(datelst))]
    
    tmelst=[]
    for time in datelst:
        try:
            tmptime=datetime.datetime.strptime(str(time).rstrip("+00:00").rstrip("+00:00"), '%Y-%m-%d %H:%M:%S')
            tmelst.append(tmptime)
        except:
            tmptime=datetime.datetime.strptime(str(time).rstrip("+00:00").rstrip("+00:00"), '%Y-%m-%d %H:%M')
            tmelst.append(tmptime)
    #        pass
    if not os.path.exists(mail_dir):
        os.mkdir(mail_dir)
          
#    mail_df=pd.DataFrame(columns=['sender_name','subject','creationtime','body','date'])
    columns=mail_df.columns
    mail_df[columns[0]]=sender_list
    mail_df[columns[1]]=msglst
    mail_df[columns[2]]=tmelst
    mail_df[columns[3]]=msgbdy
    mail_df[columns[4]]=mail_df[columns[2]].dt.date
    mail_df.to_csv(mail_dir+'\\'+str(type_)+'.csv')
    return mail_dir+'\\'+str(type_)+'.csv'
    
#    plt.scatter(mail_df['date'].value_counts().index,mail_df['date'].value_counts())
    




#def Compare_two_keywords(pytrend):
#    number =int(easygui.textbox(msg='How many keywords you want to compare ?'))
#    kw_list=[]
#    for i in range(0,number):
#        search_term=easygui.textbox(msg='Enter the Keyword '+str(i),title='Comparing  search terms')
#        kw_list.append(str(search_term))
##        search_term_2=easygui.textbox(msg='Enter the Keyword 2',title='Comparing two search terms')
#    
#    pytrend.build_payload(kw_list=kw_list)
#    document = Document()
#
#    document.add_heading('Visualizing the popularity of the keywords by using no of google search per keyword', 0)
#    p=document.add_paragraph('No of Keywords to be comapred'+str(number))
#
#
## Interest Over Time
#    interest_over_time_df = pytrend.interest_over_time()
#    for i in range(0,number):
#        plt.plot(interest_over_time_df.index,interest_over_time_df[kw_list[i]],label=str(kw_list[i]))
##        sns.tsplot(data=interest_over_time_df[str(kw_list[i])],legend=True)
#        plt.xlabel('Time')
#        plt.title('comparison of search terms')
#        plt.legend()
##        plt.
##        sns.tsplot(data=interest_over_time_df[str(kw_list[]])
##    plt.save
#    
#    plt.savefig('interest_over_time.png')
#    adding_break=p.add_run()
#    adding_break.add_break()    
#    p.add_paragraph('visualizing the popularity')    
#    adding_break=p.add_run()
#    adding_break.add_break()    
#    p.add_picture('interest_over_time.png')
#
#    
#    
#    interest_by_region_df = pytrend.interest_by_region()
#    for i in range(0,number):
#        Top_10_search_term=interest_by_region_df.sort_values(by=str(kw_list[i]))[-10:]
#        sns.barplot(Top_10_search_term[str(kw_list[i])],Top_10_search_term.index,data=interest_over_time_df)
#        plt.savefig('interest_by_region'+str(kw_list[i])+'.png')
#        p.add_paragraph('visualizing the popularity by region')    
#        p.add_picture('interest_by_region'+str(kw_list[i])+'.png')
        
#dir(easygui)\
def mailDocumentcreation(cdir,contentlist):
    try:
        document = Document()
        pngfiles=[]
        for content in contentlist:
            
            document.add_heading(text=str(content)+'Documentation')
            
            pngfiles.extend(glob.glob(cdir+'\\'+str(content)+'*.png'))
            for png in pngfiles:
                document.add_paragraph(str(png.split('\\')[-1].split('.')[0]))
                document.add_picture(png)
        document.save(cdir+'\\'+str(content)+'.docx')
    except Exception as e:
        speak.Speak('Sorry My mistake please provide your feedback regarding this error')
        easygui.exceptionbox(str(e))
        
        

    
def explanatory_analysis_timesheet(filepath,specific_day=False,specific_org=False,cont='0'):
    try:
        df=pd.read_csv(filepath,encoding = "ISO-8859-1")
        content='Timesheet'
        if specific_day:
            df['custom']=pd.DatetimeIndex(df['Start']).date
            df=df[df['custom']==pd.DatetimeIndex([cont])[0].date()]
            cont=cont.replace('/','_')
            content=content+cont
            
#            content=co
        if specific_org:
#            df=df[df['Organizer']==cont]
            df['custom']=df['Organizer'].str.lower()
            df[df['custom'].str.contains(str(cont))]
            content=content+cont
        df.index=pd.DatetimeIndex(df.Start)
        df.drop('Unnamed: 0',axis=1,inplace=True)
        org_grouped=df.groupby('Organizer')
        org_sumed=org_grouped.sum()
        org_sumed.sort_values(by='Duration(Minutes)',inplace=True)
        #plt.barh(org_sumed[-10:].index,org_sumed[-10:]['Duration(Minutes)'])
    #    fig, ax = plt.subplots()
        # the size of A4 paper
        #fig.set_size_inches(10, 5)
        print(org_sumed[-10:].index,org_sumed[-10:]['Duration(Minutes)'].values)
        sns.barplot(org_sumed[-10:]['Duration(Minutes)'].values,org_sumed[-10:].index)
        plt.tight_layout()
        print(mail_dir+'\\'+str(content)+' Timesheet organisers top10.png')
        plt.savefig(mail_dir+'\\'+str(content)+' Timesheet organisers top10.png',bbox_inches="tight")
        
        plt.close()
        
        
        
        df['date']=df.index.date
        day_wise_meeting=df.groupby('date').sum()
        print(day_wise_meeting.index,day_wise_meeting['Duration(Minutes)'])
        plt.plot(day_wise_meeting.index,day_wise_meeting['Duration(Minutes)'])
        plt.xticks(rotation=25)
        plt.xlabel('Date')
        plt.ylabel('Meetings(Minutes)')
        plt.tight_layout()
        plt.savefig(mail_dir+'\\'+str(content)+' top10 maximum date.png',bbox_inches='tight')
        
        plt.close()
        #sns.tsplot(day_wise_meeting['Duration(Minutes)'])
        
    #    title=df.Title
    #    plt.scatter(df.index,df.Title)
    #    plt.tight_layout()
        
    #    plt.savefig('testing.png',bbox_inches='tight')
        
    #    table_dict={}
    #    table_dict['windstream Demo Discussion']='windstream'
        df['lower_title']=df.Title.str.lower()
        list_title=[]   
        color_list=[]
        for string in df['lower_title'].values:
            if string.find('ML')>=0:
                list_title.append('ML')
                color_list.append(1)
            elif string.find('Devops')>=0:
                list_title.append('Devops')
                color_list.append(2)
            elif string.find('Android')>=0:
                list_title.append('Android')
                color_list.append(3)
            elif string.find('ingenious')>=0:
                list_title.append('Igenious Connect')
                color_list.append(4)
            elif string.find('rpa')>=0:
                list_title.append('RPA')
                color_list.append(5)
            elif string.find('coffee')>=0:
                list_title.append('Coffee events')
                color_list.append(6)
            elif string.find('microservice')>=0:
                list_title.append('Microservice')
                color_list.append(7)
            elif string.find('birthday')>=0:
                list_title.append('Birthday')
                color_list.append(8)
            
            else:
                list_title.append('others')
                color_list.append(9)
        df['keyword_title']=list_title
        df['color_map']=color_list
        #plt.show(bbox_inches='tight')
        df.to_pickle('prepared_data.pkl')
        plt.scatter(df.index,df.keyword_title,s=df['Duration(Minutes)'],c=df['color_map'],cmap='viridis')
        plt.xlabel('Date')
        plt.ylabel('Meeting categories')
        plt.xticks(rotation=25)
        plt.tight_layout()
        plt.savefig(mail_dir+'\\'+str(content)+' title categories scatter.png')
        plt.close()
    except Exception as e:
        speak.Speak('Sorry My mistake please provide your feedback regarding this error')
        easygui.exceptionbox(str(e))
        

def eda_mail(filepath,content,specific_day=False,specific_org=False,cont='0'):
#    pass
    try:
        idf=pd.read_csv(filepath,encoding = "ISO-8859-1")
        col=idf.columns

        if specific_day:
            idf['custom']=pd.DatetimeIndex(idf[col[-1]]).date
            idf=idf[idf['custom']==pd.DatetimeIndex([cont])[0].date()]
#            idf=idf[idf[col[-1]]==cont]
            cont=cont.replace('/','_')
            content=content+cont
        if specific_org:
            idf['custom']=idf[col[1]].str.lower()
            idf[idf['custom'].str.contains(str(cont))]
            content=content+cont

        sender_name=idf[col[1]]
        idf.index=pd.DatetimeIndex(idf[col[-1]])
        
        daywisemail=idf[col[-1]].value_counts()[:10]
        plt.bar(daywisemail.index,daywisemail.values)
        plt.xticks(rotation=25)
        plt.xlabel('Days')
        plt.ylabel('No of mails')
        plt.title('Top 10 days when you received maximum mails')
        plt.tight_layout()
        plt.savefig(mail_dir+'\\'+str(content)+'top10 daywise mails.png')
        plt.close()
#        rt=idf[col[-3]]
        
        if content=='inbox' and not specific_org:
            sender_name=idf[col[1]]
            sender_name=sender_name.value_counts()
            try:
                sender_name=sender_name.drop(['4M4U','Internal Updates','Internal Tools Support','IT Help Desk','PRODAPTIAN CONNECT','Financedesk'],axis=0)
            except:
                pass
            topmailsenders=sender_name[:10]
            plt.barh(topmailsenders.index,topmailsenders.values)
            plt.xlabel('No of mails')
            plt.ylabel('Members')
            plt.tight_layout()
            plt.savefig(mail_dir+'\\'+str(content)+'top10 senders.png')
            plt.close()
        list_title=[]
        color_list=[]
        #count_dict={}
        count_dict={i:0 for i in range(1,12)}
        idf[col[2]]=idf[col[2]].str.lower()
        
        for string in idf[col[2]].values:
            if isinstance(string,str):
                if string.find('accepted')>=0:
                    list_title.append('Meeting accepted')
                    color_list.append(10)
                    count_dict[10]+=1
                
                elif string.find('Devops')>=0:
                    list_title.append('Devops')
                    color_list.append(1)
                    count_dict[1]+=1
                elif string.find('ML')>=0:
                    list_title.append('ML')
                    color_list.append(2)
                    count_dict[2]+=1
                elif string.find('Android')>=0:
                    list_title.append('Android')
                    color_list.append(3)
                    count_dict[3]+=1
                elif string.find('ingenious')>=0:
                    list_title.append('Igenious Connect')
                    color_list.append(4)
                    count_dict[4]+=1
                elif string.find('Web development')>=0:
                    list_title.append('Web development')
                    color_list.append(5)
                    count_dict[5]+=1
                elif string.find('RPA')>=0:
                    list_title.append('RPA')
                    color_list.append(6)
                    count_dict[6]+=1
                elif string.find('microservice')>=0:
                    list_title.append('Microservice')
                    color_list.append(7)
                    count_dict[7]+=1
                elif string.find('missed conversation')>=0:
                    list_title.append('Missed conversation')
                    color_list.append(8)
                    count_dict[8]+=1
                
                else:
                    list_title.append('others')
                    color_list.append(9)
                    count_dict[9]+=1
            else:
                list_title.append('invalid mail')
                color_list.append(11)
                count_dict[11]+=1
        idf.index=pd.DatetimeIndex(idf[col[3]])
        plt.scatter(idf.index,list_title,s=list(count_dict.values()),c=color_list,cmap='viridis')
        plt.xlabel('Date')
        plt.ylabel('Mail categories')
        plt.xticks(rotation=25)
        plt.legend()
        #plt.colorbar()
        plt.tight_layout()
        plt.savefig(mail_dir+'\\'+str(content)+'subject categories scatter.png')
        plt.close()
        if content=='inbox':
            doc = Document()
            its=idf[idf.sender_name=='Internal Tools Support']
            its=its.drop('body',axis=1)
            its=its.drop('Unnamed: 0',axis=1)
            # add a table to the end and create a reference variable
            # extra row is so we can add the header row
            t = doc.add_table(its.shape[0]+1, its.shape[1])
            
            # add the header rows.
            for j in range(its.shape[-1]):
                t.cell(0,j).text = its.columns[j]
            
            # add the rest of the data frame
            for i in range(its.shape[0]):
                for j in range(its.shape[-1]):
                    t.cell(i+1,j).text = str(its.values[i,j])
            
            # save the doc
            doc.save(mail_dir+'\Internal Tools Support.docx')
        else:
            pass
    except Exception as e:
         speak.Speak('Sorry My mistake please provide your feedback regarding this error')
         easygui.exceptionbox(str(e))



if __name__=='__main__':

    speak.Speak("Hi, I am your Assistant, what can I do for you")
    intro_message = "Hi I am your Assistant, what can I do for you"
    easygui.msgbox(intro_message,title='I am your assitant',ok_button='Hello')
    cap_msg="I can help you with the following task choose the one you want"
    
    cap_list=['Timesheet','Meeting time analysis','Mail analysis','Summarize an article(give url)','How I spent the Day','Popularity estimator','Journals Downloading','Feedback']
    response_=easygui.choicebox(cap_msg,title='MY capabilities',choices=cap_list)
    
    if response_=='Timesheet':
        _=timesheet()
    if response_=='Meeting time analysis':
        filepath=timesheet()
        explanatory_analysis_timesheet(filepath)
        mailDocumentcreation(mail_dir,['Timesheet'])
    if response_=='Mail analysis':
        fi=get_mail('inbox')
        eda_mail(fi,'inbox')
        print(fi)
        fs=get_mail('sentbox')
        print(fs)
        eda_mail(fs,'sentbox')
        mailDocumentcreation(mail_dir,['inbox','sentbox'])
    #    eda_mail()
        
    if response_=='How I spent the Day':
        try:
            speak.Speak('By which you want to query your mail')
            resp=easygui.choicebox(msg='By which you want to query your mail',choices=['Organiser/sender','date(mm/dd/yyyy)'])
            if resp=='Organiser/sender':
                filepath=timesheet()
                speak.Speak('Enter the organiser or sender name')
                org=str(easygui.textbox(msg='Enter the organiser/sender name',title='Query the mail')).lower()
                explanatory_analysis_timesheet(filepath,specific_org=True,cont=org)
                
        #        explanatory_analysis_timesheet(filepath)
                fi=get_mail('inbox')
                eda_mail(fi,'inbox',specific_org=True,cont=org)
    #            fs=get_mail('sentbox')
    #            eda_mail(fs,'sentbox',specific_org=True,cont=org)
                mailDocumentcreation(mail_dir,['Timesheet'+str(org),'inbox'+str(org)])
            if resp=='date(dd/mm/yyyy)':
                filepath=timesheet()
                speak.Speak('Enter the specific date')
                date=str(easygui.textbox(msg='Enter the date(mm/dd/yyyy)',title='Query the mail'))
                explanatory_analysis_timesheet(filepath,specific_day=True,cont=date)
                
        #        explanatory_analysis_timesheet(filepath)
                fi=get_mail('inbox')
                eda_mail(fi,'inbox',specific_day=True,cont=date)
                fs=get_mail('sentbox')
                eda_mail(fs,'sentbox',specific_day=True,cont=date)
                mailDocumentcreation(mail_dir,['Timesheet'+str(date),'inbox'+str(date),'sentbox'+str(date)])
        except Exception as e:
            speak.Speak('Sorry My mistake please provide your feedback regarding this error')
            easygui.exceptionbox(str(e))
    
        #    timesheet_with_graph_for_specificDay()
    
    if response_=='Summarize an article(give url)':
        LANGUAGE = "english"
        SENTENCES_COUNT=easygui.choicebox(msg='Enter the number of sentences after summarization',choices=[10,20,30])
        summarize(SENTENCES_COUNT)
    if response_=='Journals Downloading':
        journal_downloading()
    if response_=='Popularity estimator':
        google_popularity()
    #if response_=='Reddit latest':
    #    pass
    if response_=='amazon-purchase':
        amazon_products()
    if response_=='Hackernoon news':
        pass
    if response_=='Feedback':
        feedback_mail()
    
   
#def google_trends():
#from pytrends.request import TrendReq
#import seaborn as sns
#sns.set_style('whitegrid')
#import datetime
# Login to Google. Only need to run this once, the rest of requests will use the same session.
        
        
#        Top_10_search_term=interest_by_region_df.sort_values(by='pizza')[-10:]
    
    


    
#if not glob.glob(mail_dir+'\inbox.csv'):
#    p
#else :
#    os.remove(mail_dir+'\inbox.csv')
#
    
 
    
    
        
    
    
    
    
    
#with open('df_html.html','wb') as file:
#    file.write(str(df.to_html()))
#n=df.keyword_title
#
#for i, txt in enumerate(n):
#    plt.annotate(txt, (df.index[i],df.keyword_title[i]))


#plt.ylabel('Organizeers')
#plt.xlabel('Duration in minutes')
#ax=plt.gcf()
#ax.show()

#plt.savefig()
# plt.xticks(rotation=35)


#index_counts=df.index.value_counts()
#index_counts.sort_index(inplace=True)
#plt.plot(df.index,df.Title)
#plt.plot(index_counts.index,index_counts)
##plt.xlabel('y',rotation=90)
#title=df.Title.value_counts()
#plt.bar(title[:10].index,title[:10])    
#plt.axis(rotation=34)     
#    SENTENCES_COUNT = 10
#if reponse_=='Text to Speech convertion':
#    engine = CreateObject("SAPI.SpVoice")
#    stream = CreateObject("SAPI.SpFileStream")
#
#    infile = "c:/temp/text.txt"
#    outfile = "c:/temp/text4.wav"
#    stream.Open(outfile, SpeechLib.SSFMCreateForWrite)
#    engine.AudioOutputStream = stream
#
#    f = open(infile, 'r')
#    theText = f.read()
#    f.close()
#
#    engine.speak(theText)
#    
#    stream.Close()
    
#E:/chromedriver.exe
        
        
#import requests
#from bs4 import BeautifulSoup
#
## Get login form
#URL = 'https://www.linkedin.com/uas/login'
#session = requests.session()
#login_response = session.get('https://www.linkedin.com/uas/login')
#login = BeautifulSoup(login_response.text)
#
## Get hidden form inputs
#inputs = login.find('form', {'name': 'login'}).findAll('input', {'type': ['hidden', 'submit']})
#
## Create POST data
#post = {input.get('name'): input.get('value') for input in inputs}
#post['session_key'] = 'gurunath'
#post['session_password'] = 'password'
#
## Post login
#post_response = session.post('https://www.linkedin.com/uas/login-submit', data=post)
#
## Get home page
#home_response = session.get('http://www.linkedin.com/nhome')
#home = BeautifulSoup(home_response.text)
#        
#        
#        
#import urllib
#urlopener= urllib.request.build_opener()
#urlopener.addheaders = [('User-agent', 'Mozilla/5.0')]
#html= urlopener.open('https://www.linkedin.com/in/rajagurunath//').read()
#
#from bs4 import BeautifulSoup
#from selenium import webdriver
##driver = webdriver.Firefox(executable_path='E:\Firefox_Setup_57.0.1_en-US.exe')
#driver=webdriver.Chrome('E:/chromedriver.exe')
#profile_link="https://www.linkedin.com/in/rajagurunath"
#driver.get(profile_link)
#html=driver.page_source
#soup=BeautifulSoup(html) #specify parser or it will auto-select for you
#summary=soup.find('section', { "id" : "summary" })
#print( summary.getText())
#
#
#import webbrowser
#webbrowser.open('http://www.google.com')
#
#
#
#from selenium.webdriver.common.desired_capabilities import DesiredCapabilities
#
#cap = DesiredCapabilities().FIREFOX
#cap["marionette"] = False
#browser = webdriver.Firefox(capabilities=cap, executable_path="E:\Firefox_Setup_57.0.1_en-US.exe")
#browser.get('http://google.com/')
#browser.quit()
#
#
#firefox_capabilities = DesiredCapabilities.FIREFOX
#firefox_capabilities['marionette'] = True
#firefox_capabilities['binary'] = '/usr/bin/firefox'
#browser = webdriver.Firefox(capabilities=firefox_capabilities)
#
#
#from selenium import webdriver
#from selenium.webdriver.firefox.firefox_binary import FirefoxBinary
#
#binary = FirefoxBinary('E:\Firefox_Setup_57.0.1_en-US.exe')
#browser = webdriver.Firefox(firefox_binary=binary)
#
#
#
#








#title = ""
#
#if easygui.boolbox(message, title, ["She loves me", "She loves me not"]):
#    msgbox("True ", ok_button="Good job!")
##    easygui.sendher("Flowers") # This is just a sample function that you might write.  
#
#else:
#    pass
#msg ="What is your favorite flavor?"
#title = "Ice Cream Survey"
#choices = ["Vanilla", "Chocolate", "Strawberry", "Rocky Road"]
#choice = choicebox(msg, title, choices)
#
#msgbox("you have choosen "+str(choice), ok_button="Good job!")
#
#
#
#"""
#summarization 
#Text to speech
#"""
#
#
#
#from sumy.parsers.html import HtmlParser
#from sumy.parsers.plaintext import PlaintextParser
#from sumy.nlp.tokenizers import Tokenizer
#from sumy.summarizers.lsa import LsaSummarizer as Summarizer
#from sumy.nlp.stemmers import Stemmer
#from sumy.utils import get_stop_words


#LANGUAGE = "english"
#SENTENCES_COUNT = 10
#

#if __name__ == "__main__":
##    url = "https://machinelearningmastery.com/prepare-text-data-machine-learning-scikit-learn/"
#    url="https://medium.com/@gurunathrajagopal/what-happens-when-machines-are-curious-to-learn-9aed6805bf36"
#    parser = HtmlParser.from_url(url, Tokenizer(LANGUAGE))
#    # or for plain text files
#    # parser = PlaintextParser.from_file("document.txt", Tokenizer(LANGUAGE))
#    stemmer = Stemmer(LANGUAGE)
#
#    summarizer = Summarizer(stemmer)
#    summarizer.stop_words = get_stop_words(LANGUAGE)
#    string_dict={}
#    for idx,sentence in enumerate(summarizer(parser.document, SENTENCES_COUNT)):
##        f.write(str(sentence))
#        string_dict[idx]=str(sentence)
##        print(type(sentence))
#    print(string_dict)
#    f=open(r'C:\Users\gurunath.lv\Desktop\outlook_calendar\some_text1.txt','w') 
#    f.write('Summarization')
#    for idx,sent in zip(string_dict.keys(),string_dict.values()):
#        f.write('\n')
#        f.write(str(idx))
#        f.write('.  ')
#        f.write(sent)
#    f.close()

#    f_read=open(r'C:\Users\gurunath.lv\Desktop\outlook_calendar\some_text.txt','r') 
#
#from gtts import gTTS
#import os
#tts = gTTS(text=string_, lang='en')
#tts.save(r"D:\dummies\frontend\templates\good.mp3")
#os.system("mpg321 good.mp3")        
#
#
#infile = "D:\\dummies\\frontend\\templates\\some_text.txt"
#
#engine = CreateObject("SAPI.SpVoice")
#
#f = open(infile, 'r')
#theText = f.read()
#f.close()
#
#engine.speak(theText)

#from comtypes.client import CreateObject


#
#import PyPDF2
#
#

